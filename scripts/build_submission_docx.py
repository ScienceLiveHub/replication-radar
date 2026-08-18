#!/usr/bin/env python
# Fill the official OpenAIRE hackathon .docx template with the Replication Radar submission.
import re, copy
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.table import Table
from docx.text.paragraph import Paragraph

import os
# Regenerates the submission .docx from the official template + the content in this file.
# Usage:  pip install python-docx  &&  python scripts/build_submission_docx.py
# The full submission text lives inline below; edit here and re-run to rebuild the .docx.
REPO = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
SRC = os.path.join(REPO, "OpenAIRE_AI_Hackathon_Submission_Template_final.docx")
OUT = os.path.join(REPO, "OpenAIRE_AI_Hackathon_Submission_ReplicationRadar.docx")
IMG = os.path.join(REPO, "docs", "architecture.png")
LIFECYCLE = os.path.join(REPO, "docs", "research-lifecycle.png")
doc = Document(SRC)

# ---------- inline markdown -> runs (**bold**, *italic*, `code`, [t](u)) ----------
TOKEN = re.compile(r'(\*\*.+?\*\*|\*.+?\*|`.+?`|\[[^\]]+\]\([^)]+\))')
def set_runs(par, text, size=None):
    for part in TOKEN.split(text):
        if not part:
            continue
        b=i=code=False; t=part
        if part.startswith('**') and part.endswith('**'): b=True; t=part[2:-2]
        elif part.startswith('*') and part.endswith('*'): i=True; t=part[1:-1]
        elif part.startswith('`') and part.endswith('`'): code=True; t=part[1:-1]
        elif part.startswith('['):
            m=re.match(r'\[([^\]]+)\]\(([^)]+)\)', part); t=m.group(1) if m else part
        r=par.add_run(t); r.bold=b; r.italic=i
        if code: r.font.name='Consolas'
        if size: r.font.size=Pt(size)

def new_par_after(ref_p, style=None):
    p=OxmlElement('w:p'); ref_p._p.addnext(p)
    np=Paragraph(p, ref_p._parent)
    if style:
        try: np.style=doc.styles[style]
        except KeyError: pass
    return np

def find_par(prefix):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix):
            return p
    raise KeyError(prefix)

def empties_after(anchor):
    """empty (or whitespace) paragraphs immediately following anchor, until a non-empty."""
    out=[]; el=anchor._p.getnext()
    while el is not None and el.tag==qn('w:p'):
        par=Paragraph(el, anchor._parent)
        if par.text.strip(): break
        out.append(par); el=el.getnext()
    return out

def fill_after(prefix, blocks, drop_empties=True):
    """Insert blocks=[(style,text),...] after the guidance paragraph `prefix`."""
    anchor=find_par(prefix)
    empties=empties_after(anchor) if drop_empties else []
    last=anchor
    for style,text in blocks:
        p=new_par_after(last, style)
        set_runs(p, text)
        last=p
    for e in empties:
        e._p.getparent().remove(e._p)
    return last

def cell_set(cell, text, bold_all=False):
    cell.text=""
    p=cell.paragraphs[0]
    set_runs(p, text)
    if bold_all:
        for r in p.runs: r.bold=True

def tick(cell_or_par, needle):
    """Replace the ☐ that precedes `needle` with ☑, rewriting to plain text."""
    obj=cell_or_par
    txt=obj.text
    if '☐ '+needle in txt: txt=txt.replace('☐ '+needle, '☑ '+needle)
    elif needle=='' : txt=txt.replace('☐','☑')
    else: txt=txt.replace('☐ '+needle, '☑ '+needle)
    # rewrite
    if isinstance(obj, Paragraph):
        for r in list(obj.runs): r.text=''
        if obj.runs: obj.runs[0].text=txt
        else: obj.add_run(txt)
    else:  # cell
        obj.text=txt

def tables():
    return doc.tables

# ================= 0. Submission details =================
T0=doc.tables[1]  # 7x2 details
cell_set(T0.rows[0].cells[1], "Replication Radar: Find What's Been Checked, What's Worth Replicating, and How to Replicate It")
# theme row (single cell, 4 options)
tc=T0.rows[1].cells[1]; tc.text=tc.text.replace('☐ B - Build','☑ B - Build');
cell_set(T0.rows[2].cells[1], "Science Live Team")
tp=T0.rows[3].cells[1]; tp.text=tp.text.replace('☐ Team','☑ Team')
cell_set(T0.rows[4].cells[1], "Norway (contact: VitenHub AS); distributed team: Norway (Jean), UK (Saranjeet), Spain (Anne)")
cell_set(T0.rows[5].cells[1], "Jean Iaquinta (VitenHub AS)")
cell_set(T0.rows[6].cells[1], "jiaquinta@vitenhub.no")

# team members table (5x4): header + 3 members
TM=doc.tables[2]
members=[("Anne Fouilloux","LifeWatch ERIC","Build: Radar web app + MCP, and the marine-heatwave replication","0000-0002-1784-2920 / annefou"),
         ("Jean Iaquinta","VitenHub AS","Lead & CEO; story, video and narration (voices the demo video)","0000-0002-8763-1643"),
         ("Saranjeet Kaur Bhogal","Imperial College London","RSE review & good-practice signals","0000-0002-7038-1457")]
for i,(n,a,role,oid) in enumerate(members, start=1):
    cell_set(TM.rows[i].cells[0], n); cell_set(TM.rows[i].cells[1], a)
    cell_set(TM.rows[i].cells[2], role); cell_set(TM.rows[i].cells[3], oid)

# ================= 1.1 Overall =================
p11=("Everything runs client-side against public APIs, with no login and no keys in the browser.")
fill_after("Suggested maximum: 400 words", [
 ("Normal","**The problem.** A citation has always been shorthand for *“someone checked this.”* That shorthand is breaking. When anyone can generate a plausible paper and pad it with citations to work it has nothing to do with, counting citations, or walking the citation graph, tells you what is *popular*, not what is *true*. And a paper is not one thing you can trust: it buries many claims, and what you actually rely on is one claim. The unit that matters is the claim; the question that matters is *has this claim been proven?*"),
 ("Normal","**Replication Radar** answers that on the OpenAIRE Graph. Search a field and, for each high-impact claim, it shows whether anyone has independently proven it: a signed Science Live replication verdict (validated / contested / refuted), read live from the nanopublication network, author-agnostic. Where the Graph knows only *cited / not-cited*, the Radar adds *checked / not-checked*."),
 ("Normal","If a claim has not been proven, the Radar hands you what you need to prove it: the independent, reusable software OpenAIRE holds for that field, ranked not by citations (useless for software) but by signals a one-off deposit can't fake (a resolvable repository · GitHub stars · Software Heritage archival · fair-software.eu + RSE good practice: documented · tested · CI · contributing), and, via the OpenAIRE MCP, independent data. You run the replication and publish the proof as a signed nanopublication constellation (claim → study → outcome → evidence) that the next person can cite. The proof, not the reputation."),
 ("Normal","It is **for** researchers deciding what to build on, research software engineers (whose work is finally seen and credited), meta-scientists mapping a field, and, as an MCP server run next to the OpenAIRE MCP, AI agents that need grounded, verified signals instead of guesses. Everything is grounded (every signal from a named source) and runs client-side against public APIs, with no login and no keys in the browser: a static site anyone can fork, plus a `pip install`-able MCP."),
 ("Normal","**AI assists; people do the research.** Every link in the chain is human work: the researcher who frames the claim, the fieldworkers who gather the data, the modellers, and above all the research software engineers who make software reusable. Making the chain checkable makes each of them visible and creditable."),
])

# ================= 1.2 SWOT =================
Tswot=doc.tables[3]  # 3x3: rows [hdr],[Internal S|W],[External O|T]
def cell_bullets(cell, label, bullets):
    cell.text=""
    r=cell.paragraphs[0].add_run(label); r.bold=True
    for b in bullets:
        set_runs(cell.add_paragraph(), "• "+b)
cell_bullets(Tswot.rows[1].cells[1], "Strengths", [
  "Builds on the OpenAIRE Graph and enriches it, joining live data the Graph alone doesn't hold (GitHub, Software Heritage, Zenodo, the nanopublication network)",
  "Answers two questions citations can't: has this claim been independently checked, and does its software follow good reuse practices (documented, tested, CI)?",
  "Verdicts are signed and claim-level: citable proofs, not opinions",
  "Fully grounded: every signal traces to a named public source; nothing is guessed",
  "The web app is hosted, so it opens in any browser on any OS (nothing to install); the MCP is a one-line pip / uvx install",
])
cell_bullets(Tswot.rows[1].cells[2], "Weaknesses", [
  "Finds only what an OpenAIRE keyword search surfaces, so recall is limited",
  "Shows a verdict only where a replication has been published as a nanopublication, a young and thin corpus today",
  "The software score reads reuse signals (docs, tests, CI, archival) from a public repo; these predict reusability but do not prove it, which is confirmed only when the software is actually reused",
  "Covers only two trust signals (independent replication, and software reuse practices), not other dimensions of trustworthiness such as data quality or statistical rigour; a first step, not a complete trust layer",
  "Tested on Linux; the MCP is not yet verified on Windows or macOS (the hosted web app is unaffected)",
])
cell_bullets(Tswot.rows[2].cells[1], "Opportunities", [
  "A grounded foundation AI agents can build on: verified facts instead of guesses",
  "Broaden coverage by bringing existing replication corpora (curated databases, registered reports, retraction/erratum signals) into the same open format",
  "Could capture demand by letting people flag which claims they want replicated",
])
cell_bullets(Tswot.rows[2].cells[2], "Threats", [
  "Its value grows only as more replications get published on the network",
  "Relies on external APIs (OpenAIRE, GitHub) staying stable and within rate limits",
  "Inherits quirks in OpenAIRE's own metadata (e.g. subject classifications)",
])

# ================= 1.3 The story (four parts) =================
# The question
fill_after("What did you set out to find out or solve", [
 ("Normal","When you cite a paper, what are you really citing? Not twenty pages, but **one claim** you've chosen to build on. For as long as science has had citations, your assurance that the claim is true has been that others cited it too: a citation as shorthand for *someone checked this*. That shorthand is breaking. When a plausible paper (and a wall of citations to work it has nothing to do with) can be generated on demand, counting citations, or walking the citation graph, tells you what is *popular*, not what is *true*."),
 ("Normal","So we set out to make the honest version workable on the OpenAIRE Graph: **before you cite a claim, know whether anyone has independently proven it, and if no one has, be handed exactly what you need to prove it yourself.** Truth lives at the claim, not the paper, so verification (and citation) belong there. It matters because everything downstream (the next study, the policy, the model an agent builds on) inherits the trust we place in that one claim."),
])
# The journey
fill_after("What you actually did: the steps", [
 ("Normal","We started paper-first: rank the papers in a field worth replicating. It worked, but it only re-served the Graph's single signal, impact, and when we *used* it we hit a wall: on a live topic search, essentially no paper carries a linked code or data artefact, because OpenAIRE rarely links materials to a paper. The undervalued, actionable things weren't the papers; they were the **software** and the **verdict**."),
 ("Normal","Two turns followed. First, the **verdict** layer. A replication is not a paper: it earns no citations and has no node in the Graph. But Science Live publishes replication outcomes as signed nanopublications (the FORRT chain: Quote → Claim → Study → Outcome → CiTO). We overlay them live, matched *by template, not by person* (on the nanopublication trusty hash) so any signer's verdict counts, and retraction-aware, so each claim carries its verdict, right beside the paper. Second, the **software**. A research software engineer, Saranjeet Kaur Bhogal (Imperial College London), told us what an RSE actually looks for: not a single “FAIR” badge, but whether code is *documented, tested, runs CI, invites contribution*. So beside the papers we surface the reusable software OpenAIRE already holds for a field, ranked by signals a one-off study deposit can't fake (a resolvable repository · GitHub stars · Software Heritage archival · fair-software.eu + those RSE practices). We kept everything grounded: we even deleted a feature we'd built on a guess (a keyword-matched “relevant software” picker that surfaced off-topic repos) because guessing relevance isn't anchored in any named source."),
 ("Normal","Then we closed the loop, for real. An agent running the Radar *beside* the OpenAIRE MCP took a claim the Radar flags as high-impact but **un-replicated**: Oliver et al. 2018, marine heatwaves (Nature Communications, ~1,740 citations). It pushed back on our first data choice: ERA5's sea-surface temperature shares the original paper's HadISST lineage, so it isn't independent, and it switched to independent ESA CCI satellite SST, detecting events with the independent XMHW detector, cross-checked against the author's own code. The replication is now **published as a signed constellation, and it holds: Validated.** Marine-heatwave days rose 31.8 over 1982–2016 against the paper's 30 (within 6 %), frequency and duration agree, and the trend survives ENSO removal, independent in *both* data and code. We kept the scope honest: satellite SST only reaches back to ~1982, so this validates the satellite-era trend and leaves the paper's century-scale figure explicitly out of scope. The proof, its verdict, and plain-language retellings for citizens and for schools are a citable constellation on the network, not a claim taken on trust."),
])
# The insight
fill_after("What you found, built or proved", [
 ("Normal","Truth lives at the **claim**, not the paper, so verification and citation belong there. Reliability (*was it proven?*) and reusability (*is the software good?*) are different categories of signal: you *add* them live on top of the Graph; you don't repair citation counts into them. A signed replication constellation turns a proof into a first-class, citable object, and the very same move finally makes research-software-engineering work visible and creditable, not invisible behind the paper authors' names."),
 ("Normal","And it's built to be read by humans as much as by agents: a person sees the software behind a claim and whether it's been checked; an agent, running the Radar beside the OpenAIRE MCP, gets both the structural graph and the verification-and-reuse layer, grounded data it can cite instead of hallucinate. AI assists (it finds, connects, drafts, checks) but people do the research. It's a first brick of a graph of *verified, engineered* knowledge for agentic science."),
])
# What others can reuse
fill_after("The single most important part", [
 ("Normal","• **The live web app:** pure static, queries OpenAIRE + GitHub / Software Heritage / Zenodo + the nanopublication network from the browser; fork it and point it elsewhere."),
 ("Normal","• **An MCP server** (`pip install replication-radar`) exposing the same engine to any MCP-capable agent, to run beside the OpenAIRE MCP."),
 ("Normal","• **A grounded software assessment:** fair-software.eu recommendations plus RSE good-practice signals (documented / tested / CI / contributing), computed live from GitHub + Software Heritage + Zenodo, no third-party scorer."),
 ("Normal","• **A reproducible, author-agnostic, retraction-aware verdict-index method:** FORRT Outcome / CiTO nanopublications joined on the trusty hash."),
 ("Normal","• **A machine-readable methodology & provenance spec** (`methodology.json`, CC-BY): every signal's source and formula; plus a feasibility map of which open-science APIs are reachable and CORS-friendly."),
])

# hero figure at the top of §1.3 (the story-page illustration; matches online)
introp=find_par("This is the 1")
_f=new_par_after(introp,"Normal")
_f.add_run().add_picture(LIFECYCLE, width=Inches(6.3))
_fc=new_par_after(_f,"Normal")
_rc=_fc.add_run("Figure 1. Two ways to trust a published claim: today, by reputation (citation counts, which can be manufactured); or with Replication Radar, by proof (re-run the claim with independent data and software, sign the outcome, and it becomes citable). Discovery and verdicts run live on the OpenAIRE Graph and the Science Live nanopublication network.")
_rc.italic=True; _rc.font.size=Pt(9)

# ================= 2.1 How it works =================
last21=fill_after("Suggested maximum: 300 words", [
 ("Normal","A static, client-side site (HTML/CSS/JS on Netlify) plus one serverless proxy for the GitHub / ScholeXplorer APIs that lifts the keyless rate limit and CORS while keeping any token server-side. Every signal is fetched live; the client holds no keys. The same engine ships as an MCP server (`replication-radar`, FastMCP)."),
 ("Normal","**Pipeline (see diagram).** 1) **Discover.** OpenAIRE Graph REST API: papers (type=publication, BIP! impact C1–C5 + impulse) and reusable software (type=software, ranked by grounded reusability (resolvable repo · stars · fair-software.eu + RSE practices · Software Heritage), since citation class is useless for software). 2) **Prove?** Read replication outcomes live from the Science Live nanopublication network via SPARQL: FORRT Outcome + CiTO nanopublications (author-agnostic, by template), joined on the trusty hash, retraction-filtered via the admin graph, overlaid by DOI; each verdict carries the atomic claim it tested. 3) **Assess software.** For any resolvable repo, compute fair-software.eu + RSE practices + language / paid-runtime, live from GitHub (proxied), Software Heritage, Zenodo. 4) **Connect.** The paper↔software edge from the Graph's own relations (ScholeXplorer / OpenAIRE MCP) links the two lenses. 5) **Prove & publish.** The FORRT template scaffolds a replication; the proof is signed as a nanopublication constellation that re-surfaces on the Radar for the next person."),
])
# insert the diagram image after the §2.1 text
imgp=new_par_after(last21, "Normal")
imgp.add_run().add_picture(IMG, width=Inches(6.3))
cap=new_par_after(imgp, "Normal")
r=cap.add_run("Figure 2. Replication Radar architecture. Two entry points (researcher on the web app, an AI agent on the MCP beside the OpenAIRE MCP) read the OpenAIRE Graph, Science Live nanopublications, and GitHub / Software Heritage / Zenodo into a ranked queue with verdicts and software signals; proving a claim publishes a signed constellation back to the network, closing the loop.")
r.italic=True; r.font.size=Pt(9)

# ================= 2.2 Graph elements table =================
T22=doc.tables[4]  # 7x2
det={
 1:"Anonymous, CORS-enabled REST. Papers: `GET api.openaire.eu/graph/v1/researchProducts?search=<topic>&type=publication&sortBy=influence DESC` (paged). Reusable software: the same call with `type=software`. A claim's DOI is resolved to its OpenAIRE record via the dedup id `?id=doi_dedup___::md5(lowercased-doi)&pageSize=1`. Funder context: `GET /projects?search=<topic>`.",
 2:"OpenAIRE MCP (Alien gateway, OAuth), run beside our own `replication-radar` MCP: `explore_research_relationships` for the paper↔software / dataset edges (e.g. `isSourceOf`), and a dataset lookup our `find_dataset` tool hands off to, to find an independent replication dataset with a citable DOI.",
 3:"`researchProducts` of type publication, software and dataset; `projects` (funders).",
 4:"`indicators.citationImpact.{influenceClass, citationClass, impulseClass, citationCount}` (BIP! classes C1–C5); `type`; `pids` (DOI); `subjects` (FOS, SDG); `openAccessColor`; relation types (`isSourceOf`, `isSupplementedBy`, `cites`).",
 5:"Science Live nanopublication network (SPARQL `query.knowledgepixels.com/repo/full` + admin graph; read for FORRT Outcome `hasValidationStatus` / `hasOutcomeRepository` and CiTO verdict relations, matched to OpenAIRE products by DOI); GitHub REST API; Software Heritage API; Zenodo API; ScholeXplorer; Copernicus Marine (ESA CCI SST, for the replication).",
 6:"Impact-ranked over the full matching set per topic. Live product counts: `search=species distribution model` returns 61,284; `machine learning climate` 27,300; `bumblebee decline` 297. Verdicts read live from the nanopublication network (the set grows as replications publish); software assessed per resolvable repo (GitHub 60/hr unauthenticated, cached).",
}
for i,txt in det.items():
    cell_set(T22.rows[i].cells[1], txt)

# ================= 2.3 Documentation =================
fill_after("Suggested maximum: 250 words", [
 ("Normal","**Web app.** Nothing to install: open https://openaire-hackathon.netlify.app. Locally: clone and serve `site/` with any static server (`python3 -m http.server`); pure HTML/CSS/JS, no build."),
 ("Normal","**MCP server.** `pip install replication-radar`, register with any MCP client (`{\"mcpServers\":{\"replication-radar\":{\"command\":\"replication-radar\"}}}`). Tools: `radar`, `replication_status`, `find_independent_software`, `verified_claims`: public sources only."),
 ("Normal","**Reproduce the demo.** A step-by-step walkthrough (exact searches, prompts, results) at https://openaire-hackathon.netlify.app/demo.html."),
 ("Normal","**Provenance.** Every label/score is documented signal-by-signal in `methodology.json` (machine-readable, CC-BY), rendered at /methodology.html; /replicate.html explains the end-to-end prove-then-cite loop."),
 ("Normal","**Repo.** README + architecture; MIT LICENSE; CITATION.cff; deps in `pyproject.toml`; full history; no credentials. Archived on Zenodo every release: concept DOI https://doi.org/10.5281/zenodo.21850976 (latest v0.4.6)."),
])

# ================= 3. Innovation & risks =================
fill_after("What does this do that existing tools", [
 ("Normal","Most research-quality tools try to improve the metric. The Radar instead treats the claim, not the paper, as the unit of trust, and adds two categories the citation graph structurally can't hold: has this claim been independently checked, and does its software follow good reuse practices? Both are read live on top of the Graph. The verdict layer is claim-level, cryptographically-signed nanopublications read author-agnostically by template, so an LLM can cite it instead of hallucinating. Three moves are, to our knowledge, novel: (1) making OpenAIRE's software assessable, not just findable; (2) using the Graph's own relations to link a paper to its software (Soroye 2020, `isSourceOf`); and (3) packaging it as an MCP beside the OpenAIRE MCP, a first brick of verified knowledge for agentic science."),
])
fill_after("Where does it break", [
 ("Normal","The biggest limit is verdict coverage: today the overlay reflects only replications already published as nanopublications, a young and thin corpus, so many claims read as “not yet checked”. By design the overlay is source-agnostic (any signer, matched by template), so it grows as replications publish, and as existing replication corpora are minted into the same open format. Beyond that: discovery recall is keyword-bound; software assessment needs a public repo, and GitHub's 60/hr unauthenticated limit caps throughput (cached). A green CI run is a reproducibility signal, not independent reproduction, which the verdict covers separately. We add two of the missing layers, not the whole graph."),
])
fill_after("Apart from the Alien AI MCP Connector", [
 ("Normal","The artifact is deterministic: no LLM in the data path, nothing hallucinated on screen. AI enters two ways. (1) The deliverable is an MCP connector for any MCP-capable agent, run beside the OpenAIRE MCP, not tied to one model or vendor. (2) Development used Claude Code (Anthropic) throughout, including this write-up, always under human direction and review, with every data source verified against its live API before shipping. Human-in-the-loop by design: AI assists; people decide, do the research, and are credited."),
])
fill_after("Confirm no personal data was processed", [
 ("Normal","No personal data is processed: only public scholarly metadata (OpenAIRE, the nanopublication network, GitHub / Software Heritage / Zenodo); no accounts or private data stored (the client holds no keys; the GitHub-proxy token is server-side). Third-party data and code are used within licence: OpenAIRE and the nanopub network are open; ESA CCI SST is cited by DOI; reused software is open-source, credited by repo. Materials: CC-BY 4.0 (write-up), MIT (code)."),
])

# ================= 4. Links & artifacts table =================
def add_hyperlink(par, url, text):
    part=par.part; r_id=part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    h=OxmlElement('w:hyperlink'); h.set(qn('r:id'), r_id)
    run=OxmlElement('w:r'); rpr=OxmlElement('w:rPr')
    c=OxmlElement('w:color'); c.set(qn('w:val'),'0563C1'); rpr.append(c)
    u=OxmlElement('w:u'); u.set(qn('w:val'),'single'); rpr.append(u)
    run.append(rpr); t=OxmlElement('w:t'); t.text=text; run.append(t); h.append(run); par._p.append(h)
    return h

T4=doc.tables[5]  # 9x4: Item, Link, Status, Notes
rows4={
 1:("https://github.com/ScienceLiveHub/replication-radar","☑ Public","MIT"),
 2:("https://openaire-hackathon.netlify.app","☑ Public","Static, client-side"),
 3:("https://youtu.be/hVyLafY3Y3E","☑ Public","2:49"),
 4:("https://openaire-hackathon.netlify.app","☑ Public","The Replication Radar tool (opens on a live sample scan, as in the video); also a pip-installable MCP (pip install replication-radar)"),
 5:("https://openaire-hackathon.netlify.app/replicate.html","☑ Public","+ reproduce-the-demo walkthrough /demo.html"),
 6:("https://openaire-hackathon.netlify.app/story.html","☑ Public","The 1.3 story (CC-BY)"),
 7:("https://doi.org/10.5281/zenodo.21850976","","Concept DOI (latest v0.4.6)"),
 8:("https://annefou.github.io/marine-heatwave-replication/","","Example proof: independent replication of Oliver et al. 2018 (marine heatwaves), Validated in both data and code. Reproducible study (Jupyter Book); code + data at github.com/annefou/marine-heatwave-replication (Zenodo 10.5281/zenodo.21950032); signed constellation on the nanopublication network."),
}
for i,(url,status,notes) in rows4.items():
    linkcell=T4.rows[i].cells[1]; linkcell.text=""
    add_hyperlink(linkcell.paragraphs[0], url, url)
    if status: cell_set(T4.rows[i].cells[2], status)
    cell_set(T4.rows[i].cells[3], notes)

# full signed nanopub list for the marine-heatwave example proof, inserted after the §4 table
NANOPUBS = [
 ("Quote + personal comment", "https://w3id.org/sciencelive/np/RA1m-2tHCFBVjypflopbfGnuEpxrzPE_khGhss1FJMtbA"),
 ("AIDA sentence", "https://w3id.org/sciencelive/np/RAmbschSgs8k_AoM34DgIWT0EMkrwKYKEWQFIXkslMtfM"),
 ("FORRT Claim", "https://w3id.org/sciencelive/np/RAGw-EZjva3ybpqWtY3loRFToZrm6GSMQqamvoe-e-JjE"),
 ("FORRT Replication Study", "https://w3id.org/sciencelive/np/RAdIP7v2kJyOD-hRDIKdZkjMfFHkOIKdWoSnvAeGasU_s"),
 ("FORRT Replication Outcome (Validated; the story apex)", "https://w3id.org/sciencelive/np/RAGjvtR-Pq6576AIEsj5CTiLW3yK0cPEbfgk7OjhxyVVM"),
 ("CiTO Citation (confirms Oliver 2018)", "https://w3id.org/sciencelive/np/RAnns3mUVRk1WNb6SAuAunHiGGZL4eD9c_Axx-8eCj0x4"),
 ("Story (the constellation)", "https://w3id.org/sciencelive/np/RAQfGMNmJiFt4KDDJE8dsiT3az4CRYehRuDFoa-tXGc8w"),
 ("Plain-language summary for the public", "https://w3id.org/np/RAwGVnm2oxnLVNYR18dmkFqbqFL4ZYw3yX0x0CxXzdzS8"),
 ("Plain-language summary for schools", "https://w3id.org/np/RAaYmBJ1IAyn75wkl_G3zDC7NmV7Rz2KlwU6J3f5JQujY"),
]
_anchor = T4._tbl
_hp = OxmlElement('w:p'); _anchor.addnext(_hp); _hpar = Paragraph(_hp, doc)
_hr = _hpar.add_run("Signed nanopublications (the full marine-heatwave constellation, all published live on the nanopublication network):"); _hr.bold = True; _hr.font.size = Pt(10)
_last = _hp
for _label, _url in NANOPUBS:
    _p = OxmlElement('w:p'); _last.addnext(_p); _par = Paragraph(_p, doc)
    _r = _par.add_run("• " + _label + ": "); _r.font.size = Pt(9)
    add_hyperlink(_par, _url, _url)
    _last = _p

# repository checklist (paragraphs starting with ☐)
for pref in ["☐   README explains","☐   LICENSE file present","☐   Dependencies listed",
             "☐   Commit history visible","☐   No credentials"]:
    try: p=find_par(pref); tick(p,'')
    except KeyError: pass

# ================= 5. Openness & licensing =================
T5=doc.tables[6]  # 6x2
T5.rows[0].cells[1].text=T5.rows[0].cells[1].text.replace('☐ confirmed','☑ confirmed')
cell_set(T5.rows[1].cells[1], "MIT")
cell_set(T5.rows[2].cells[1], "CC-BY 4.0")
T5.rows[3].cells[1].text=T5.rows[3].cells[1].text.replace('☐ confirmed','☑ confirmed')
T5.rows[4].cells[1].text=T5.rows[4].cells[1].text.replace('☐ confirmed','☑ confirmed')
T5.rows[5].cells[1].text=T5.rows[5].cells[1].text.replace('☐ confirmed','☑ confirmed')

# ================= 7. Final check (tick the done ones) =================
for pref in ["☐   Theme selected","☐   Story (section 1.3) written","☐   Video under 3 minutes",
             "☐   CC-BY 4.0 applied","☐   Contact email correct"]:
    try: p=find_par(pref); tick(p,'')
    except KeyError: pass
# leave: "Every link ... private browser window" and "Submitted before ..." for Anne to confirm

# Remove the template's §3 guidance prompts + the suggested-max note for a clean, unambiguous section
def del_par(prefix):
    for p in list(doc.paragraphs):
        if p.text.strip().startswith(prefix):
            p._p.getparent().remove(p._p); return True
    return False
for pref in ["What does this do that existing",
             "Where does it break, what did you not",
             "Apart from the Alien AI MCP Connector",
             "Confirm no personal data was processed",
             "Suggested maximum: 400 words in total"]:
    del_par(pref)

doc.save(OUT)
print("SAVED:", OUT)
